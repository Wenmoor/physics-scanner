from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from config import settings
from models.content import ContentSegment, ExtractedProblem
from services.omml_writer import append_inline_math

FONT_CN = "宋体"
FONT_SIZE = Pt(10.5)  # 五号
LINE_SPACING = Pt(15.75)  # 五号单倍行距
PROBLEM_GAP_LINES = 1  # 多题之间的间隔行数


def _page_line_capacity() -> int:
    content_height_mm = (
        settings.page_height_mm - settings.margin_top_mm - settings.margin_bottom_mm
    )
    content_height_pt = content_height_mm * 2.834645669
    return max(1, int(content_height_pt / LINE_SPACING.pt))


def _count_content_lines(problem: ExtractedProblem) -> int:
    lines = len(problem.paragraphs) or 1
    math_count = sum(1 for para in problem.paragraphs for seg in para if seg.type == "math")
    return lines + (math_count + 1) // 2


def _distribute_blank_lines(problems: list[ExtractedProblem]) -> list[int]:
    """在一页内按预估权重分配手写留白行数。"""
    page_lines = _page_line_capacity()
    content_total = sum(_count_content_lines(p) for p in problems)
    gap_total = PROBLEM_GAP_LINES * max(0, len(problems) - 1)
    remaining = page_lines - content_total - gap_total

    if len(problems) == 1:
        return [max(0, remaining)]

    if remaining <= 0:
        return [1] * len(problems)

    weights = [max(p.handwriting_lines, 1) for p in problems]
    weight_sum = sum(weights)
    blanks = [max(1, int(remaining * w / weight_sum)) for w in weights]

    diff = remaining - sum(blanks)
    idx = 0
    while diff != 0 and blanks:
        if diff > 0:
            blanks[idx % len(blanks)] += 1
            diff -= 1
        elif blanks[idx % len(blanks)] > 1:
            blanks[idx % len(blanks)] -= 1
            diff += 1
        idx += 1
        if idx > len(blanks) * 40:
            break

    return blanks


def _set_run_font(run) -> None:
    run.font.size = FONT_SIZE
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def _set_paragraph_format(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def _configure_b5_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(settings.page_width_mm)
    section.page_height = Mm(settings.page_height_mm)
    section.left_margin = Mm(settings.margin_left_mm)
    section.right_margin = Mm(settings.margin_right_mm)
    section.top_margin = Mm(settings.margin_top_mm)
    section.bottom_margin = Mm(settings.margin_bottom_mm)


def _add_segments_paragraph(doc: Document, segments: list[ContentSegment]) -> None:
    p = doc.add_paragraph()
    _set_paragraph_format(p)
    for seg in segments:
        if seg.type == "text":
            run = p.add_run(seg.value)
            _set_run_font(run)
        elif seg.type == "math":
            try:
                append_inline_math(p, seg.value)
            except Exception:
                run = p.add_run(seg.value)
                _set_run_font(run)


def _add_blank_lines(doc: Document, count: int) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        _set_paragraph_format(p)


def _add_problem_content(doc: Document, problem: ExtractedProblem) -> None:
    for para_segments in problem.paragraphs:
        _add_segments_paragraph(doc, para_segments)


def build_docx(
    problems: list[ExtractedProblem],
    output_path: Path | None = None,
) -> Path:
    if not problems:
        raise ValueError("没有可排版的题目")

    settings.output_dir.mkdir(parents=True, exist_ok=True)
    if output_path is None:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = settings.output_dir / f"{stamp}_错题本.docx"

    blank_lines = _distribute_blank_lines(problems)

    doc = Document()
    _configure_b5_page(doc)

    for idx, problem in enumerate(problems):
        if idx > 0:
            _add_blank_lines(doc, PROBLEM_GAP_LINES)
        _add_problem_content(doc, problem)
        _add_blank_lines(doc, blank_lines[idx])

    doc.save(output_path)
    return output_path
